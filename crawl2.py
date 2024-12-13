from bs4 import BeautifulSoup
import docx
import urllib.request
import time
import json
import requests
import os
import os.path
userAgents=[
'Mozilla/5.0 (Linux; Android 10; K) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Mobile Safari/537.36',
'Mozilla/5.0 (Linux; Android 10; K) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36',
'Mozilla/5.0 (Linux; Android 10; K) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Mobile Safari/537.36,gzip(gfe)'
]
csname='KH01'
doc = docx.Document()
story = 1
#with open('.\cookies.json') as f:#
    #data = json.load(f)#
#cookie = random.choice(data['cookies'])#
cookie="_ga=GA1.1.620714396.1733953857; truyenyy-cookies-consent=accepted; messages=.eJyLjlaKj88qzs-Lz00tLk5MT1XSMdAxMtVRCi5NTgaKpJXm5FQqFGem56WmKGTmKSQWKyTm5GcnWhroKcXqDE3NsQDztlK9:1tLfoU:PP_LmixhaFGolb9sk-Ngy62xWxlZq9ZP_LMeBFVRbks; csrftoken=ZI5A6GU0WCOiTAO4i7citNAlEH1RJnkLtIMT8LeZncRv2Vn4LVCja6jjvuKLnj7J; truyenyyid=gnc70xdiwjxcd4xxpex8lzjtpdv2k4j5; cf_clearance=A6_GTfcMmwZRQpdHEWWIXullE2AhApwONwNXA01eCzc-1734008116-1.2.1.1-n24wSlsMsmx7ocCaY42huyidMspuaBWwbpVOT3DakqH_JJcCeBPX_oF2nsU7k3keJiSuiri9adHfyT81Mj9NZiuAX4Ok6H2vgOdy8PDHjgDq7vKy7NQjn0ZuL7ylyIKBsEvBy1WWo1FwtXHjltSo3GBB62DTKMOWboZF_kmvoQM5rMH3FjoVIcIQ9hnjo2hn2ZHYCcuENMrymbZMhfgPg35rylbdZDyxodpeQkvFrNq897nya3Z_ckqS4pZiFtlT7pFE26gwLvWq5ycroc4n_4lN2JwpPdotAXsWjS4eXZHRfltHoAH9ubwO5X9rqVG68KAZm8BFKCJNivVk4pIQ4FHPeEVbLit6T.44OYfuHrG_36i2tt_odLUgUGPQNx.Sa5qcVBgm4GJjGGNr589g_g; _ga_6SE80VQKJ2=GS1.1.1734007797.6.1.1734008254.0.0.0"
start=0
end=0
domain = 'https://truyenyy.vip'

def get_html(page,name):
    list_link=[]
    url_html =  'https://truyenyy.vip/truyen/'+name+'/danh-sach-chuong/?p='+str(page)
    request_html = urllib.request.Request(url_html)
    request_html.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36 Edg/131.0.0.0')

    page_html = urllib.request.urlopen(request_html)
    
    soup = BeautifulSoup(page_html, 'html.parser')

    tbody = soup.find('table', class_='table').find('tbody')
    a_elements = tbody.find_all('a',class_='table-chap-title')
    for a in a_elements:
        list_link.append({'link':a.get('href'),'title': a.get_text()})
    return list_link

def get_content(link, file_name):
    list_content=[]
    url_content =  domain+link['link']
    request_content = urllib.request.Request(url_content)
    request_content.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36 Edg/131.0.0.0')
    request_content.add_header('Cookie', cookie)

    page_content = urllib.request.urlopen(request_content)
    time.sleep(5)
    soup = BeautifulSoup(page_content, 'lxml')
    story_name=link['title']
    if soup.find('div',{"id": "vip-content-placeholder"}):
        id_chap = soup.find('div',{"id": "jq-dropdown-1"}).find('ul',{"class": "jq-dropdown-menu"}).find('li').find('a').get('href').replace('/account/settings/?chap=','')
        for i in range(0,5):
            request_chap = urllib.request.Request('https://truyenyy.vip/web-api/novel/chapter-content-get/?chap_id='+id_chap+'&part='+str(i))
            request_chap.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36 Edg/131.0.0.0')
            request_chap.add_header('Cookie', cookie)
            page_chap = urllib.request.urlopen(request_chap)
            data = page_chap.read().decode('utf-8')
            json_obj = json.loads(data)
            if json_obj['content'] != '':
                soup_chap = BeautifulSoup(json_obj['content'], 'lxml')
                for sentence in soup_chap.find_all('p'):
                    for child in sentence.find_all(recursive=True):
                        if child.has_attr("style"):
                            child.decompose()
                        if child.name == 'style':
                            child.decompose()
                    soup_chap.prettify()
                    list_content.append(sentence.get_text()) 
            else:
                break 
    else:
       p_elements = soup.find('div',{"id": "inner_chap_content_1"}).find_all('p') 
       for p in p_elements:
           list_content.append(p.get_text(strip=True))
    write_Word(list_content, story_name, file_name)
    return list_content

def write_Word(content, name, file_name):
    doc = docx.Document()
    check_file = os.path.isfile(".\{}.docx".format(file_name))
    if check_file:
        f = open('{}.docx'.format(file_name), 'rb')
        doc = docx.Document(f)
    p = doc.add_paragraph()
    p.add_run('Chương '+str(story)+': '+ name + '\n').bold = True
    for c in content:
        run = p.add_run(c)
        run.add_break()
    doc.save(".\{}.docx".format(file_name))

def processing(start, end, name, file_name):
    global story
    print('Khởi động chương trình...')
    pagging=['1']
    url =  domain + '/truyen/'+name+'/danh-sach-chuong/'
    request = urllib.request.Request(url)
    request.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36 Edg/131.0.0.0')
    request.add_header('Cookie',cookie)
    page = urllib.request.urlopen(request)
    soup = BeautifulSoup(page, 'html.parser')
    new_feed = soup.find('ul', class_='pagination')
    li_elements = new_feed.find_all('li',class_="page-item")
    for item in li_elements:
        t = item.find_all('a', class_="page-link")
        for aa in t:
            pagging.append(aa.text)
    
    pagging = [item for item in pagging if item.isdigit()]
    end_story_crawl = end
    try:
        for i in range(int(pagging[0]), int(pagging[-1])+1):
            list_html = get_html(i, name)
            for page_link in list_html:
                if end_story_crawl < 0:
                    return
                if start > 0:
                    start = start-1
                if start == 0:
                    print(page_link)
                    get_content(page_link, file_name)
                    story = story + 1
                    end_story_crawl = end - story
    except Exception as e:
        print('Lỗi dịch vụ, xin vui lòng liên hệ email hoangvanhieusw@gmail.com')
        print(e) 

def checkLicense():
    url = 'http://171.244.143.235:2895/weatherforecast?name='+csname+'-0'
    resp = requests.get(url=url)
    data = resp.json()
    print('Hiện tại bạn đang còn {} lượt sử dụng'.format(data))
    return data
def updateLicense():
    url = 'http://171.244.143.235:2895/weatherforecast?name='+csname+'-1'
    requests.get(url=url)

def start():
    global story
    if int(checkLicense()) <= 0:
        print('Hết hạn sử dụng! Vui lòng liên hệ email hoangvanhieusw@gmail.com')
        input("----------------------------------------------------------------")
    else:
        story_name = input('Tên truyện (Link url): ')
        question = input('Bạn có muốn in tất cả? (Y/N)')
        if question == 'Y' or question == 'y':
           updateLicense()
           processing(0,0,story_name)
        else:
           story_start = input('Bắt đầu từ Chương:')
           story_end = input('Kết thúc từ chương:')
           file_name = input('Đặt tên truyện. Lưu ý không chứa các kí tự đặc biệt: (\/:*?"<>|)')
           story = int(story_start)
           updateLicense()
           processing(int(story_start),int(story_end),story_name, file_name)
        print('Hoàn Tất!')  
if __name__=='__main__':
    start() 
    
