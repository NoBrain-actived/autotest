from bs4 import BeautifulSoup
import docx
import urllib.request
import time
import json
import random
userAgents=[
'Mozilla/5.0 (Linux; Android 10; K) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Mobile Safari/537.36',
'Mozilla/5.0 (Linux; Android 10; K) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36',
'Mozilla/5.0 (Linux; Android 10; K) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Mobile Safari/537.36,gzip(gfe)'
]
doc = docx.Document()
story = 1
with open('.\cookies.json') as f:
    data = json.load(f)
cookie = random.choice(data['cookies'])
start=0
end=0
domain = 'https://truyenyy.vip'

def get_html(page):
    list_link=[]
    url_html =  'https://truyenyy.vip/truyen/xung-ba-vo-dao/danh-sach-chuong/?p='+str(page)
    request_html = urllib.request.Request(url_html)
    request_html.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36 Edg/131.0.0.0')

    page_html = urllib.request.urlopen(request_html)
    
    soup = BeautifulSoup(page_html, 'html.parser')

    tbody = soup.find('table', class_='table').find('tbody')
    a_elements = tbody.find_all('a',class_='table-chap-title')
    for a in a_elements:
        list_link.append({'link':a.get('href'),'title': a.get_text()})
    return list_link

def get_content(link):
    list_content=[]
    url_content =  domain+link['link']
    request_content = urllib.request.Request(url_content)
    request_content.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36 Edg/131.0.0.0')
    request_content.add_header('Cookie', cookie)

    page_content = urllib.request.urlopen(request_content)
    time.sleep(5)
    soup = BeautifulSoup(page_content, 'lxml')
    story_name=link['title']
    if soup.find('div',{"id": "vip-content-placeholder"}):
        id_chap = soup.find('div',{"id": "jq-dropdown-1"}).find('ul',{"class": "jq-dropdown-menu"}).find('li').find('a').get('href').replace('/account/settings/?chap=','')
        for i in range(1,5):
            request_chap = urllib.request.Request('https://truyenyy.vip/web-api/novel/chapter-content-get/?chap_id='+id_chap+'&part='+str(i))
            request_chap.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36 Edg/131.0.0.0')
            request_chap.add_header('Cookie', cookie)
            page_chap = urllib.request.urlopen(request_chap)
            data = page_chap.read().decode('utf-8')
            json_obj = json.loads(data)
            if json_obj['content'] != '':
                soup_chap = BeautifulSoup(json_obj['content'], 'lxml')
                for sentence in soup_chap.find_all('p'):
                    for child in sentence.find_all(recursive=True):
                        if child.has_attr("style"):
                            child.decompose()
                        if child.name == 'style':
                            child.decompose()
                    soup_chap.prettify()
                    list_content.append(sentence.get_text()) 
            else:
                break 
    else:
       p_elements = soup.find('div',{"id": "inner_chap_content_1"}).find_all('p') 
       for p in p_elements:
           list_content.append(p.get_text(strip=True))
    write_Word(list_content, story_name)
    return list_content

def write_Word(content, name):
    p = doc.add_paragraph()
    p.add_run('Chương '+str(story)+': '+ name + '\n').bold = True
    for c in content:
        run = p.add_run(c)
        run.add_break()
    doc.save("D:/PythonDocument/demo.docx")

def processing(start, end, name):
    global story
    print('Khởi động chương trình...')
    pagging=['1']
    url =  domain + '/truyen/'+name+'/danh-sach-chuong/'
    request = urllib.request.Request(url)
    request.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36 Edg/131.0.0.0')
    request.add_header('Cookie',cookie)
    page = urllib.request.urlopen(request)
    soup = BeautifulSoup(page, 'html.parser')
    new_feed = soup.find('ul', class_='pagination')
    li_elements = new_feed.find_all('li',class_="page-item")
    for item in li_elements:
        t = item.find_all('a', class_="page-link")
        for aa in t:
            pagging.append(aa.text)
    
    pagging = [item for item in pagging if item.isdigit()]
    end_story_crawl = end
    try:
        for i in range(int(pagging[0]), int(pagging[-1])+1):
            list_html = get_html(i)
            for page_link in list_html:
                if end_story_crawl < 0:
                    return
                if start > 0:
                    start = start-1
                if start == 0:
                    print(page_link)
                    get_content(page_link)
                    story = story + 1
                    end_story_crawl = end - story
    except Exception as e:
        print('Lỗi dịch vụ, xin vui lòng liên hệ email hoangvanhieusw@gmail.com')
        print(e) 

if __name__=='__main__':
    story_name = input('Tên truyện (Link url): ')
    question = input('Bạn có muốn in tất cả? (Y/N)')
    if question == 'Y':
       processing(0,0,story_name)
    else:
       story_start = input('Bắt đầu từ Chương:')
       story_end = input('Kết thúc từ chương:')
       story = int(story_start)
       processing(int(story_start),int(story_end),story_name)
    print('Hoàn Tất!')      
    
